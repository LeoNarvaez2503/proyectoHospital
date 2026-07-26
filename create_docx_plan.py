import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def build_docx_document(input_txt_path, output_docx_path):
    doc = docx.Document()
    
    # Configuración de Márgenes (1 pulgada)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Encabezado Superior Derecho (Times New Roman)
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("proyectoHospital SQA Plan\nTM-SQA-01 v2.0\n25 de Julio de 2026")
        hrun.font.name = 'Times New Roman'
        hrun.font.size = Pt(9)
        hrun.font.color.rgb = RGBColor(0, 0, 0)
        
    # Estilos Globales: FUENTE TIMES NEW ROMAN
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0, 0, 0)
    
    # ---------------------------------------------------------------------------
    # PÁGINA 1: PORTADA OFICIAL DE LA PLANTILLA
    # ---------------------------------------------------------------------------
    p_top = doc.add_paragraph()
    p_top.paragraph_format.space_before = Pt(72)
    p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_top.add_run("SOFTWARE QUALITY ASSURANCE PLAN\nTEMPLATE\n\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(18)
    r.bold = True
    
    r_code = p_top.add_run("TM-SQA-01 V2.0\n\n\n\n\n\n\n\n")
    r_code.font.name = 'Times New Roman'
    r_code.font.size = Pt(16)
    r_code.bold = True
    
    p_bot = doc.add_paragraph()
    p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bot.paragraph_format.space_before = Pt(120)
    r_org = p_bot.add_run(
        "Systems Engineering Process Office, Code 212\n"
        "Space and Naval Warfare Systems Center San Diego\n"
        "53560 Hull Street\n"
        "San Diego, CA 92152-5001\n\n\n"
        "Approved for public release; distribution is unlimited"
    )
    r_org.font.name = 'Times New Roman'
    r_org.font.size = Pt(11)
    r_org.bold = True
    
    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # PÁGINA 2: DOCUMENT CHANGE REQUEST (DCR)
    # ---------------------------------------------------------------------------
    p_dcr = doc.add_paragraph()
    p_dcr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_dcr = p_dcr.add_run("DOCUMENT CHANGE REQUEST (DCR)")
    r_dcr.font.name = 'Times New Roman'
    r_dcr.font.size = Pt(14)
    r_dcr.bold = True
    p_dcr.paragraph_format.space_after = Pt(12)
    
    t_dcr = doc.add_table(rows=7, cols=2)
    t_dcr.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_dcr.style = 'Table Grid'
    
    dcr_data = [
        ("Document Title: Software Quality Assurance Plan Template", "Tracking Number: TM-SQA-01 v2.0"),
        ("Name of Submitting Organization: Departamento de Ingeniería de Software y Calidad", ""),
        ("Organization Contact: Equipo SQA proyectoHospital", "Phone: +593 99 625 9054"),
        ("Mailing Address: Quito, Ecuador", ""),
        ("Short Title: proyectoHospital SQAP", "Date: 25 de Julio de 2026"),
        ("Change Location:\n(use section #, figure #, table #, etc.)\n\nProposed Change:\nDocumentación consolidada del Plan Maestro de Aseguramiento de Calidad con 71 pruebas automatizadas xUnit, 12 escenarios de Postman, prueba de estrés masivo de 50 usuarios / 1,200 inserciones BD, auditoría Nielsen SUS 85/100 y remediación de 16 defectos Jira.", ""),
        ("Rational for Change:\nCumplimiento de estándares de calidad IEEE 730-2014, ISO/IEC 25010 y OWASP ASVS v4.0.", "")
    ]
    
    for row_idx, (c1, c2) in enumerate(dcr_data):
        row = t_dcr.rows[row_idx]
        cell1 = row.cells[0]
        cell2 = row.cells[1]
        
        cell1.paragraphs[0].text = c1
        cell2.paragraphs[0].text = c2
        
        for c in (cell1, cell2):
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            
    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(8)
    r_n = p_note.add_run(
        "Note: For the Systems Engineering Process Office (SEPO) to take appropriate action on a change request, "
        "please provide a clear description of the recommended change along with supporting rationale.\n"
        "Send to: Commanding Officer, Space and Naval Warfare Systems Center, Code 212, 53560 Hull Street, San Diego, CA 92152-5001\n"
        "Fax: (619) 553-6249 | Email: sepo@spawar.navy.mil | Submit online: http://sepo.spawar.navy.mil/\n"
    )
    r_n.font.name = 'Times New Roman'
    r_n.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # PÁGINA 3: DOCUMENT CONVENTIONS
    # ---------------------------------------------------------------------------
    p_conv = doc.add_paragraph()
    p_conv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c = p_conv.add_run("DOCUMENT CONVENTIONS")
    r_c.font.name = 'Times New Roman'
    r_c.font.size = Pt(14)
    r_c.bold = True
    p_conv.paragraph_format.space_after = Pt(12)
    
    conv_text = (
        "This document is a Software Quality Assurance (SQA) Plan template. As such, wording in this document "
        "should be supplemented with project-specific information to produce an SQA Plan that accurately "
        "describes the project SQA organization. Therefore, tailor (add, delete, change, or expand) the information "
        "provided in this document.\n\n"
        "Standard conventions are used within this document to direct the reader to specific sections of the text. "
        "These sections provide instructions and explanations and require users to substitute their own department-specific "
        "information for the generic information provided or to \"fill in a blank.\"\n\n"
        "Regular text\tStandard SQA plan content. Tailor regular text to accurately describe the project SQA organization.\n\n"
        "In some cases where information may already be found in another project document, like the Software Development "
        "Plan (SDP), refer to that document rather than duplicate the information in the project SQA Plan.\n\n"
        "The template begins with the Project SQA cover sheet on the page after the next. Update the header page to reflect "
        "the document configuration identifier for the project SQA Plan."
    )
    p_ct = doc.add_paragraph()
    r_ct = p_ct.add_run(conv_text)
    r_ct.font.name = 'Times New Roman'
    r_ct.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # PÁGINA 4: PROJECT COVER SHEET & APPROVALS
    # ---------------------------------------------------------------------------
    p_proj = doc.add_paragraph()
    p_proj.paragraph_format.space_before = Pt(48)
    p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_pr = p_proj.add_run("proyectoHospital\nSOFTWARE QUALITY ASSURANCE PLAN\n\n")
    r_pr.font.name = 'Times New Roman'
    r_pr.font.size = Pt(18)
    r_pr.bold = True
    
    r_pr_code = p_proj.add_run("TM-SQA-01 v2.0\n25 de Julio de 2026\n\n\n")
    r_pr_code.font.name = 'Times New Roman'
    r_pr_code.font.size = Pt(14)
    r_pr_code.bold = True
    
    p_app = doc.add_paragraph()
    p_app.paragraph_format.space_before = Pt(36)
    r_ap_title = p_app.add_run("SQA Plan Approvals:\n\n\n")
    r_ap_title.font.name = 'Times New Roman'
    r_ap_title.font.size = Pt(12)
    r_ap_title.bold = True
    
    t_app = doc.add_table(rows=3, cols=2)
    t_app.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    app_rows = [
        ("____________________________\nSQA Manager", "____________________\nDate: 25/07/2026"),
        ("____________________________\nProject Manager", "____________________\nDate: 25/07/2026"),
        ("____________________________\nProgram Manager", "____________________\nDate: 25/07/2026")
    ]
    for idx, (m1, m2) in enumerate(app_rows):
        row = t_app.rows[idx]
        row.cells[0].paragraphs[0].text = m1
        row.cells[1].paragraphs[0].text = m2
        for c in row.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run.bold = True
                    
    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # PROCESAMIENTO DE LAS 14 SECCIONES TM-SQA-01 V2.0
    # ---------------------------------------------------------------------------
    with open(input_txt_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_lines = []
    
    for line in lines:
        raw_line = line.strip()
        
        if raw_line.startswith("========"):
            if in_table and table_lines:
                process_table(doc, table_lines)
                in_table = False
                table_lines = []
            continue
            
        if '|' in raw_line and not raw_line.startswith("-----"):
            in_table = True
            table_lines.append(raw_line)
            continue
        elif in_table and (raw_line.startswith("-----") or not raw_line):
            if raw_line.startswith("-----"):
                continue
            process_table(doc, table_lines)
            in_table = False
            table_lines = []
            continue
        elif in_table and '|' not in raw_line:
            process_table(doc, table_lines)
            in_table = False
            table_lines = []
            
        if not raw_line:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            continue
            
        # Encabezados Principales (1. PURPOSE, 2. REFERENCE DOCUMENTS, etc.)
        if (len(raw_line) > 2 and raw_line[0].isdigit() and (raw_line[1] == '.' or raw_line[2] == '.')) or raw_line.startswith("TABLA DE CONTENIDOS") or raw_line.startswith("SOFTWARE QUALITY ASSURANCE PLAN"):
            h = doc.add_heading(level=1)
            run = h.add_run(raw_line)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13.5)
            run.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            continue
            
        # Subencabezados (1.1, 3.1, 7.1, etc.)
        if (len(raw_line) > 3 and raw_line[0].isdigit() and raw_line[2] == '.'):
            h = doc.add_heading(level=2)
            run = h.add_run(raw_line)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            continue
            
        # Viñetas o Listas
        if raw_line.startswith("* ") or raw_line.startswith("• ") or raw_line.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            text_clean = raw_line[2:].strip()
            format_inline_text(p, text_clean)
            p.paragraph_format.space_after = Pt(3)
            continue
            
        # Párrafo Estándar
        p = doc.add_paragraph()
        format_inline_text(p, raw_line)
        p.paragraph_format.space_after = Pt(6)
        
    if in_table and table_lines:
        process_table(doc, table_lines)
        
    doc.save(output_docx_path)
    print(f"📄 Documento Word .docx (14 Secciones TM-SQA-01 V2.0) generado en: {output_docx_path}")

def format_inline_text(paragraph, text):
    if ":" in text and not text.startswith("http"):
        parts = text.split(":", 1)
        run_bold = paragraph.add_run(parts[0] + ":")
        run_bold.bold = True
        run_bold.font.name = 'Times New Roman'
        run_bold.font.color.rgb = RGBColor(0, 0, 0)
        
        run_val = paragraph.add_run(parts[1])
        run_val.font.name = 'Times New Roman'
        run_val.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)

def process_table(doc, lines):
    if not lines:
        return
        
    matrix = []
    for line in lines:
        cells = [c.strip() for c in line.split('|')]
        if any(cells):
            matrix.append(cells)
            
    if not matrix:
        return
        
    num_rows = len(matrix)
    num_cols = max(len(row) for row in matrix)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    for i, row in enumerate(matrix):
        for j, val in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = val
                set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                
                if i == 0:
                    set_cell_background(cell, "EAEAEA")
                    
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    if i == 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(9.5)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        if i == 0 or j == 0:
                            run.bold = True
                            
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(6)

if __name__ == "__main__":
    txt_source = "/home/meatpuppets/Escritorio/University/proyectoHospital/SQA_Plan_Template.docx.txt"
    docx_target1 = "/home/meatpuppets/Escritorio/University/proyectoHospital/SQA_Plan_Template.docx"
    docx_target2 = "/home/meatpuppets/Escritorio/University/proyectoHospital/SQAP/PLAN_ASEGURAMIENTO_CALIDAD_SQAP.docx"
    
    build_docx_document(txt_source, docx_target1)
    build_docx_document(txt_source, docx_target2)
